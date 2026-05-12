#!/opt/venv/bin/python3
"""docx_inspect — structural inspection of .docx documents.

Backed by python-docx for paragraph/table/section traversal; stdlib
zipfile + ElementTree for style and field extraction where python-docx
is awkward.
"""

from __future__ import annotations

import argparse
import os
import sys
import zipfile
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple
from xml.etree import ElementTree as ET

import ooxml
import render
from docx import Document
from docx.document import Document as DocumentType
from utils import (
    TEXT_PREVIEW_LIMIT,
    ellipsize,
    format_size,
    pad,
    safe_output,
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS = ooxml.NS["a"]

DEFAULT_MAX_PARAGRAPHS = 200

USAGE = (
    "docx_inspect <file> [--styles] [--paragraphs] [--text] [--tables] "
    "[--sections] [--fields] [--media] [--render] "
    "[--offset N] [--max N] [--page N]"
)

HELP_TEXT = (
    "docx_inspect - Inspect .docx document structure\n"
    "\n"
    f"Usage: {USAGE}\n"
    "\n"
    "Options:\n"
    "  --styles      Named paragraph and character styles with direct\n"
    "                typography (typeface, size, bold, color, alignment)\n"
    "                collected along each style's basedOn chain.\n"
    "  --paragraphs  Walk paragraphs (1-indexed) with pStyle and text.\n"
    "  --text        Extract readable text with markdown markers\n"
    "                (# heading, - bullet, | table cell |).\n"
    "  --tables      Per-table rows x cols, widths, row-1 preview.\n"
    "  --sections    Page size, margins, orientation per section.\n"
    "  --fields      Fields (TOC, REF, HYPERLINK, DATE, ...) with stale flag.\n"
    "  --media       Embedded media under word/media/ with sizes.\n"
    "  --render      Rasterize pages to JPEG (100 dpi) via soffice + pdftoppm.\n"
    "                Outputs to /tmp/docx_render/<doc>/page-NNN.jpg.\n"
    "  --max N       Max paragraphs printed in --paragraphs (default 200).\n"
    "  --offset N    Skip first N paragraphs in --paragraphs (default 0).\n"
    "  --page N      Render only the given page (1-indexed) with --render.\n"
    "\n"
    "Default (no flag): document overview — counts + heading outline."
)


# ---------------------------------------------------------------------------
# Helpers.
# ---------------------------------------------------------------------------


def _qw(local: str) -> str:
    return f"{{{W_NS}}}{local}"


def _w(elem, attr: str) -> Optional[str]:
    if elem is None:
        return None
    return elem.attrib.get(_qw(attr))


def _int(text: Optional[str]) -> Optional[int]:
    if text is None:
        return None
    try:
        return int(text)
    except ValueError:
        return None


# ---------------------------------------------------------------------------
# Styles. Walk basedOn chain collecting the first direct rPr/pPr value for
# each attribute. No theme resolution, no docDefaults merge — just what
# the named style chain explicitly sets.
# ---------------------------------------------------------------------------


def _styles_xml(zf: zipfile.ZipFile) -> Optional[ET.Element]:
    return ooxml.read_xml(zf, "word/styles.xml")


def _style_index(styles_xml) -> Dict[str, ET.Element]:
    if styles_xml is None:
        return {}
    out: Dict[str, ET.Element] = {}
    for style in styles_xml.findall(_qw("style")):
        sid = style.attrib.get(_qw("styleId"))
        if sid:
            out[sid] = style
    return out


def _style_name(style) -> str:
    name_elem = style.find(_qw("name"))
    if name_elem is not None:
        val = name_elem.attrib.get(_qw("val"))
        if val:
            return val
    return style.attrib.get(_qw("styleId")) or "?"


def _fill_typography(rpr, ppr, target: Dict[str, Optional[str]]) -> None:
    """Fill in `target` from an rPr/pPr pair — only fields still None.
    Earlier (more specific) calls win over later (basedOn ancestor) calls."""
    if rpr is not None:
        rfonts = rpr.find(_qw("rFonts"))
        if rfonts is not None and target["typeface"] is None:
            face = rfonts.attrib.get(_qw("ascii"))
            if face:
                target["typeface"] = face
        sz = rpr.find(_qw("sz"))
        if sz is not None and target["size_pt"] is None:
            half_pt = _int(sz.attrib.get(_qw("val")))
            if half_pt is not None:
                target["size_pt"] = half_pt / 2.0
        b = rpr.find(_qw("b"))
        if b is not None and target["bold"] is None:
            target["bold"] = b.attrib.get(_qw("val")) not in ("0", "false")
        i = rpr.find(_qw("i"))
        if i is not None and target["italic"] is None:
            target["italic"] = i.attrib.get(_qw("val")) not in ("0", "false")
        color = rpr.find(_qw("color"))
        if color is not None and target["color"] is None:
            val = color.attrib.get(_qw("val"))
            if val and val != "auto":
                target["color"] = val.upper()
    if ppr is not None:
        jc = ppr.find(_qw("jc"))
        if jc is not None and target["align"] is None:
            val = jc.attrib.get(_qw("val"))
            if val:
                target["align"] = val


def resolve_style_typography(
    style_id: str,
    style_idx: Dict[str, ET.Element],
) -> Dict[str, Optional[str]]:
    acc: Dict[str, Optional[str]] = {
        "typeface": None,
        "size_pt": None,
        "bold": None,
        "italic": None,
        "color": None,
        "align": None,
    }
    seen = set()
    current = style_id
    while current and current not in seen:
        seen.add(current)
        style = style_idx.get(current)
        if style is None:
            break
        _fill_typography(style.find(_qw("rPr")), style.find(_qw("pPr")), acc)
        based = style.find(_qw("basedOn"))
        current = _w(based, "val") if based is not None else None
    return acc


def format_typography(t: Dict[str, Optional[str]]) -> str:
    parts: List[str] = []
    if t["typeface"]:
        parts.append(str(t["typeface"]))
    if t["size_pt"] is not None:
        size_pt = float(t["size_pt"])
        parts.append(
            f"{int(size_pt)}pt" if size_pt.is_integer() else f"{size_pt:.1f}pt"
        )
    if t["bold"]:
        parts.append("bold")
    if t["italic"]:
        parts.append("italic")
    if t["color"]:
        parts.append(f"#{t['color']}")
    if t["align"]:
        parts.append(f"algn={t['align']}")
    return " ".join(parts)


def _outline_level_for_style(
    style_id: Optional[str],
    style_idx: Dict[str, ET.Element],
) -> Optional[int]:
    """Walk basedOn looking for the first <w:outlineLvl>. This is how Word
    tags heading styles and how the TOC picks them up."""
    if not style_id:
        return None
    seen = set()
    current: Optional[str] = style_id
    while current and current not in seen:
        seen.add(current)
        style = style_idx.get(current)
        if style is None:
            break
        outline = style.find(f"{_qw('pPr')}/{_qw('outlineLvl')}")
        if outline is not None:
            return _int(_w(outline, "val"))
        based = style.find(_qw("basedOn"))
        current = _w(based, "val") if based is not None else None
    return None


# ---------------------------------------------------------------------------
# Body iteration. Walk top-level paragraphs and tables in document order.
# ---------------------------------------------------------------------------


def _body_elements(doc_xml) -> Iterable[Tuple[str, ET.Element]]:
    if doc_xml is None:
        return
    body = doc_xml.find(_qw("body"))
    if body is None:
        return
    for child in body:
        tag = child.tag.split("}", 1)[-1]
        if tag in ("p", "tbl"):
            yield tag, child


def _paragraph_text(p_xml) -> str:
    chunks: List[str] = []
    for t in p_xml.iter(_qw("t")):
        if t.text:
            chunks.append(t.text)
    for _ in p_xml.iter(_qw("br")):
        chunks.append("\n")
    return "".join(chunks).strip().replace("\n", " ")


def _paragraph_pstyle(p_xml) -> Optional[str]:
    pstyle = p_xml.find(f"{_qw('pPr')}/{_qw('pStyle')}")
    return _w(pstyle, "val") if pstyle is not None else None


# ---------------------------------------------------------------------------
# Views.
# ---------------------------------------------------------------------------


def _summarize_counts(doc_xml) -> Dict[str, int]:
    counts = {
        "paragraphs": 0,
        "tables": 0,
        "images": 0,
        "fields_simple": 0,
        "fields_complex": 0,
        "hyperlinks": 0,
    }
    if doc_xml is None:
        return counts
    for kind, _ in _body_elements(doc_xml):
        if kind == "p":
            counts["paragraphs"] += 1
        elif kind == "tbl":
            counts["tables"] += 1
    counts["images"] = sum(1 for _ in doc_xml.iter(f"{{{A_NS}}}blip"))
    counts["fields_simple"] = sum(1 for _ in doc_xml.iter(_qw("fldSimple")))
    counts["fields_complex"] = sum(
        1
        for fc in doc_xml.iter(_qw("fldChar"))
        if _w(fc, "fldCharType") == "begin"
    )
    counts["hyperlinks"] = sum(1 for _ in doc_xml.iter(_qw("hyperlink")))
    return counts


def print_overview(doc: DocumentType, doc_xml, styles_xml) -> str:
    lines: List[str] = []
    sections = doc.sections
    width_in = sections[0].page_width.inches if sections else 0
    height_in = sections[0].page_height.inches if sections else 0
    counts = _summarize_counts(doc_xml)
    lines.append(
        f"[Paragraphs: {counts['paragraphs']} | "
        f"tables: {counts['tables']} | "
        f"sections: {len(sections)} | "
        f"images: {counts['images']} | "
        f"size: {width_in:.1f}x{height_in:.1f}\"]"
    )

    extra: List[str] = []
    if counts["fields_simple"] or counts["fields_complex"]:
        extra.append(
            f"fields: {counts['fields_simple'] + counts['fields_complex']}"
        )
    if counts["hyperlinks"]:
        extra.append(f"hyperlinks: {counts['hyperlinks']}")
    if extra:
        lines.append(f"[{' | '.join(extra)}]")

    style_idx = _style_index(styles_xml)
    headings: List[Tuple[int, int, str]] = []
    para_idx = 0
    for kind, elem in _body_elements(doc_xml):
        if kind != "p":
            continue
        para_idx += 1
        level = _outline_level_for_style(_paragraph_pstyle(elem), style_idx)
        if level is None:
            continue
        text = _paragraph_text(elem)
        if not text:
            continue
        headings.append((para_idx, level, text))

    if headings:
        lines.append("")
        lines.append("[Outline]")
        for p_idx, level, text in headings:
            indent = "  " * (level + 1)
            lines.append(
                f"  #{pad(str(p_idx), 4)} {indent}H{level + 1}: "
                f"{ellipsize(text, TEXT_PREVIEW_LIMIT)}"
            )
    return "\n".join(lines)


def print_styles(styles_xml) -> str:
    if styles_xml is None:
        return "[No styles part]"
    style_idx = _style_index(styles_xml)

    paragraph_styles: List[Tuple[str, str]] = []
    character_styles: List[Tuple[str, str]] = []
    table_styles: List[Tuple[str, str]] = []
    for sid, style in style_idx.items():
        stype = style.attrib.get(_qw("type")) or "paragraph"
        name = _style_name(style)
        if stype == "character":
            character_styles.append((name, sid))
        elif stype == "table":
            table_styles.append((name, sid))
        else:
            paragraph_styles.append((name, sid))

    lines = [
        f"[Styles: {len(paragraph_styles)} paragraph, "
        f"{len(character_styles)} character, "
        f"{len(table_styles)} table]"
    ]

    def emit(group_name: str, group: List[Tuple[str, str]]) -> None:
        if not group:
            return
        group.sort(key=lambda t: t[0].lower())
        lines.append("")
        lines.append(f"# {group_name}")
        for name, sid in group:
            typo_str = format_typography(resolve_style_typography(sid, style_idx))
            head = f"- {pad(name, 28)} ({sid})"
            if typo_str:
                head += f"  {typo_str}"
            lines.append(head)

    emit("Paragraph styles", paragraph_styles)
    emit("Character styles", character_styles)
    emit("Table styles", table_styles)
    return "\n".join(lines)


def print_paragraphs(doc_xml, offset: int, max_count: int) -> str:
    if doc_xml is None:
        return "[No document body]"
    paragraphs: List[ET.Element] = [
        elem for kind, elem in _body_elements(doc_xml) if kind == "p"
    ]
    total = len(paragraphs)
    end = min(total, offset + max_count)
    visible = paragraphs[offset:end]

    lines = [f"[Paragraphs: {total} | showing {offset + 1}-{end}]"]
    for i, p in enumerate(visible, start=offset + 1):
        pstyle = _paragraph_pstyle(p) or "(default)"
        text = _paragraph_text(p)
        if not text and pstyle == "(default)":
            continue
        head = f"#{pad(str(i), 4)} pStyle={pad(pstyle, 18)}"
        if text:
            head += f' "{ellipsize(text, TEXT_PREVIEW_LIMIT)}"'
        lines.append(head)

    if end < total:
        lines.append("")
        lines.append(
            f"[Showing paragraphs {offset + 1}-{end} of {total}. "
            f"Use --offset {end} for the next page.]"
        )
    return "\n".join(lines)


def print_text(doc_xml, styles_xml) -> str:
    if doc_xml is None:
        return "[No document body]"
    style_idx = _style_index(styles_xml)
    blocks: List[str] = []
    total_chars = 0
    for kind, elem in _body_elements(doc_xml):
        if kind == "p":
            text = _paragraph_text(elem)
            if not text:
                continue
            pstyle = _paragraph_pstyle(elem)
            level = _outline_level_for_style(pstyle, style_idx)
            if level is not None:
                line = "#" * (level + 1) + " " + ellipsize(text, TEXT_PREVIEW_LIMIT)
            elif pstyle and "List" in pstyle:
                line = "- " + ellipsize(text, TEXT_PREVIEW_LIMIT)
            else:
                line = ellipsize(text, TEXT_PREVIEW_LIMIT)
            blocks.append(line)
            total_chars += len(line)
        elif kind == "tbl":
            for row in elem.findall(_qw("tr")):
                cells: List[str] = []
                for cell in row.findall(_qw("tc")):
                    cell_text = "".join(
                        (t.text or "") for t in cell.iter(_qw("t"))
                    ).strip().replace("\n", " ")
                    cells.append(ellipsize(cell_text, 40))
                blocks.append("| " + " | ".join(cells) + " |")
            blocks.append("")
    if not blocks:
        return "[No text in document]"
    while blocks and not blocks[-1]:
        blocks.pop()
    return f"[Text: {total_chars} chars]\n\n" + "\n".join(blocks)


def print_tables(doc: DocumentType) -> str:
    tables = doc.tables
    if not tables:
        return "[No tables]"
    lines = [f"[Tables: {len(tables)}]"]
    for ti, table in enumerate(tables, start=1):
        rows = table.rows
        nrows = len(rows)
        ncols = len(table.columns)
        widths_in: List[float] = []
        for col in table.columns:
            try:
                widths_in.append(col.width.inches if col.width else 0.0)
            except (AttributeError, TypeError):
                widths_in.append(0.0)
        total_in = sum(widths_in)
        widths_str = ", ".join(f"{w:.2f}\"" for w in widths_in)
        lines.append(
            f"  #{ti}  {nrows}x{ncols}  total: {total_in:.2f}\"  "
            f"cols: [{widths_str}]"
        )
        if nrows:
            cells = [
                ellipsize((c.text or "").strip().replace("\n", " "), 30)
                for c in rows[0].cells
            ]
            lines.append("      row1: | " + " | ".join(cells) + " |")
    return "\n".join(lines)


def print_sections(doc: DocumentType) -> str:
    sections = doc.sections
    if not sections:
        return "[No sections]"
    lines = [f"[Sections: {len(sections)}]"]
    for si, sec in enumerate(sections, start=1):
        try:
            w = sec.page_width.inches if sec.page_width else 0
            h = sec.page_height.inches if sec.page_height else 0
            mt = sec.top_margin.inches if sec.top_margin else 0
            mr = sec.right_margin.inches if sec.right_margin else 0
            mb = sec.bottom_margin.inches if sec.bottom_margin else 0
            ml = sec.left_margin.inches if sec.left_margin else 0
        except (AttributeError, TypeError):
            w = h = mt = mr = mb = ml = 0
        if sec.orientation is not None:
            orient = sec.orientation.name.lower()
        else:
            orient = "landscape" if w and h and w > h else "portrait"
        lines.append(
            f"  #{si}  size: {w:.1f}x{h:.1f}\"  orient: {orient}  "
            f"margins T/R/B/L: {mt:.2f}/{mr:.2f}/{mb:.2f}/{ml:.2f}\""
        )
    return "\n".join(lines)


_FIELD_PLACEHOLDER_MARKERS = (
    "Error!",
    "Table of Contents entries not found",
    "No table of contents entries found",
    "Reference source not found",
)


def _classify_field(instr: str, result: str) -> str:
    if not result and instr:
        return "empty"
    for marker in _FIELD_PLACEHOLDER_MARKERS:
        if marker in result:
            return "stale"
    return "ok"


def print_fields(doc_xml) -> str:
    if doc_xml is None:
        return "[No document body]"
    fields: List[Tuple[str, str, str]] = []

    for fs in doc_xml.iter(_qw("fldSimple")):
        instr = (_w(fs, "instr") or "").strip()
        result = "".join((t.text or "") for t in fs.iter(_qw("t"))).strip()
        fields.append((instr, _classify_field(instr, result), result))

    for kind, p in _body_elements(doc_xml):
        if kind != "p":
            continue
        runs = list(p.iter(_qw("r")))
        i = 0
        while i < len(runs):
            r = runs[i]
            fc = r.find(_qw("fldChar"))
            if fc is not None and _w(fc, "fldCharType") == "begin":
                instr_parts: List[str] = []
                result_parts: List[str] = []
                in_result = False
                j = i + 1
                while j < len(runs):
                    rj = runs[j]
                    fcj = rj.find(_qw("fldChar"))
                    if fcj is not None:
                        ct = _w(fcj, "fldCharType")
                        if ct == "separate":
                            in_result = True
                        elif ct == "end":
                            break
                    for it in rj.findall(_qw("instrText")):
                        if it.text:
                            instr_parts.append(it.text)
                    if in_result:
                        for t in rj.findall(_qw("t")):
                            if t.text:
                                result_parts.append(t.text)
                    j += 1
                instr = "".join(instr_parts).strip()
                result = "".join(result_parts).strip().replace("\n", " ")
                fields.append((instr, _classify_field(instr, result), result))
                i = j
            i += 1

    if not fields:
        return "[No fields]"
    lines = [f"[Fields: {len(fields)}]"]
    for idx, (instr, state, result) in enumerate(fields, start=1):
        kind_word = instr.split(None, 1)[0] if instr else "?"
        lines.append(
            f"  #{idx}  {pad(kind_word, 12)} state:{pad(state, 7)}  "
            f"instr: {ellipsize(instr, 50)}"
        )
        if result:
            lines.append(
                f"      result: \"{ellipsize(result, TEXT_PREVIEW_LIMIT)}\""
            )
    return "\n".join(lines)


def print_media(file_path: str) -> str:
    try:
        zf = zipfile.ZipFile(file_path)
    except zipfile.BadZipFile:
        return "[Not a valid zip / .docx package]"
    entries: List[Tuple[str, int]] = []
    with zf:
        for info in zf.infolist():
            if info.filename.startswith("word/media/"):
                entries.append((info.filename, info.file_size))
    if not entries:
        return "[No embedded media]"
    entries.sort()
    lines = [f"[Media: {len(entries)}]"]
    for name, size in entries:
        short = name[len("word/media/"):]
        lines.append(f"- {pad(short, 32)} {format_size(size)}")
    return "\n".join(lines)


def print_render(file_path: str, page_idx: Optional[int]) -> str:
    out_dir, rendered = render.render_via_soffice(
        file_path,
        out_root=Path("/tmp/docx_render"),
        item_name="page",
        item_idx=page_idx,
    )
    plural = "" if len(rendered) == 1 else "s"
    lines = [
        f"[Rendered: {len(rendered)} page{plural} | jpeg @ 100 dpi | {out_dir}]"
    ]
    for p in rendered:
        lines.append(str(p))
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# CLI dispatch.
# ---------------------------------------------------------------------------


def main() -> int:
    parser = argparse.ArgumentParser(
        prog="docx_inspect",
        usage=USAGE,
        add_help=False,
    )
    parser.add_argument("file", nargs="?")
    parser.add_argument("--styles", action="store_true")
    parser.add_argument("--paragraphs", action="store_true")
    parser.add_argument("--text", action="store_true")
    parser.add_argument("--tables", action="store_true")
    parser.add_argument("--sections", action="store_true")
    parser.add_argument("--fields", action="store_true")
    parser.add_argument("--media", action="store_true")
    parser.add_argument("--render", action="store_true")
    parser.add_argument("--max", type=int, default=DEFAULT_MAX_PARAGRAPHS)
    parser.add_argument("--offset", type=int, default=0)
    parser.add_argument("--page", type=int)
    parser.add_argument("--help", "-h", action="store_true", dest="help_flag")

    args = parser.parse_args()

    if args.help_flag:
        sys.stdout.write(HELP_TEXT + "\n")
        return 0

    if not args.file:
        sys.stderr.write(f"Error: file is required\nUsage: {USAGE}\n")
        return 1
    if not os.path.isfile(args.file):
        sys.stderr.write(f"Error: file not found: {args.file}\n")
        return 1
    if args.max < 1:
        sys.stderr.write("Error: --max must be >= 1\n")
        return 1
    if args.offset < 0:
        sys.stderr.write("Error: --offset must be >= 0\n")
        return 1

    file_header = (
        f"[File: {os.path.basename(args.file)} | "
        f"{format_size(os.path.getsize(args.file))}]"
    )

    if args.media:
        body = print_media(args.file)
    elif args.render:
        body = print_render(args.file, args.page)
    else:
        doc = Document(args.file)
        with zipfile.ZipFile(args.file) as zf:
            doc_xml = ooxml.read_xml(zf, "word/document.xml")
            styles_xml = _styles_xml(zf)

            if args.styles:
                body = print_styles(styles_xml)
            elif args.paragraphs:
                body = print_paragraphs(doc_xml, args.offset, args.max)
            elif args.text:
                body = print_text(doc_xml, styles_xml)
            elif args.tables:
                body = print_tables(doc)
            elif args.sections:
                body = print_sections(doc)
            elif args.fields:
                body = print_fields(doc_xml)
            else:
                body = print_overview(doc, doc_xml, styles_xml)

    full = file_header + "\n" + body
    text, truncated = safe_output(full)
    sys.stdout.write(text)
    sys.stdout.write("\n")
    if truncated:
        sys.stdout.write(
            "[Output truncated; narrow with a subcommand or paginate with "
            "--offset / --max]\n"
        )
    return 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except (ValueError, KeyError) as exc:
        sys.stderr.write(f"Error: {exc}\n")
        sys.exit(1)
    except Exception as exc:
        sys.stderr.write(f"Error: {type(exc).__name__}: {exc}\n")
        sys.exit(1)
